"""
온열질환 예방 현장 관리 - 사진대지 생성기
- 카테고리: 체온측정 / 체감온도계 / 물 / 그늘/휴식 / 보냉장구
- 저장/불러오기: GitHub API (records/YYYY-MM-DD.json)
- 레이아웃: 카테고리 제목 + 표 keep_with_next (줄바꿈 방지)
"""

import io, json, math, base64, time
from datetime import datetime

import streamlit as st
from PIL import Image
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── GitHub (선택적 import)
try:
    from github import Github, GithubException
    GITHUB_AVAILABLE = True
except ImportError:
    GITHUB_AVAILABLE = False

# ────────────────────────────────────────
# 상수
# ────────────────────────────────────────
FIXED_CATEGORIES = ["체온측정", "체감온도계", "물", "그늘/휴식", "보냉장구"]
TAB_ICONS        = ["🌡", "🌡️", "💧", "🌳", "🧊"]
SITE_NAME        = "성동자이리버뷰"
DOC_TITLE        = "온열질환 예방 현장 관리"
RECORDS_DIR      = "records"          # GitHub 저장소 내 폴더

# ────────────────────────────────────────
# 페이지 설정
# ────────────────────────────────────────
st.set_page_config(page_title="온열질환 예방 현장 관리", page_icon="🌡", layout="wide")

st.markdown("""
<style>
  .main-title  { font-size:1.85rem; font-weight:800; color:#1e3a5f; margin-bottom:0; }
  .sub-title   { font-size:0.86rem; color:#777; margin-bottom:0.9rem; }
  .cat-header  { background:#1a56db; color:white; border-radius:7px;
                 padding:0.35rem 0.8rem; font-weight:700; font-size:0.9rem; margin-bottom:0.28rem; }
  .count-badge { background:white; color:#1a56db; border-radius:10px;
                 padding:1px 8px; font-size:0.72rem; font-weight:700; margin-left:0.4rem; }
  .tip         { background:#fffbea; border:1px solid #f6e05e; border-radius:6px;
                 padding:0.48rem 0.75rem; font-size:0.79rem; color:#744210; margin-bottom:0.55rem; }
  .section-hd  { font-size:0.98rem; font-weight:700; color:#2d3748;
                 border-bottom:2px solid #e2e8f0; padding-bottom:0.18rem; margin:0.7rem 0 0.45rem; }
  .rec-row     { display:flex; align-items:center; gap:8px; margin-bottom:4px; }
</style>
""", unsafe_allow_html=True)

# ────────────────────────────────────────
# 세션 초기화
# ────────────────────────────────────────
if "photos" not in st.session_state:
    st.session_state.photos    = {cat: [] for cat in FIXED_CATEGORIES}
if "work_date" not in st.session_state:
    st.session_state.work_date = datetime.today().date()


# ────────────────────────────────────────
# GitHub 헬퍼
# ────────────────────────────────────────
def _gh_repo():
    """GitHub 저장소 객체 반환. secrets 없으면 None."""
    if not GITHUB_AVAILABLE:
        return None
    try:
        token = st.secrets["github"]["token"]
        repo  = st.secrets["github"]["repo"]       # "owner/repo-name"
        return Github(token).get_repo(repo)
    except Exception:
        return None


def gh_save(date_str: str, photos: dict) -> bool:
    """사진 데이터를 JSON으로 직렬화해 GitHub에 커밋."""
    repo = _gh_repo()
    if repo is None:
        return False

    serialized = {}
    for cat, items in photos.items():
        serialized[cat] = [
            {"name": name, "data": base64.b64encode(data).decode()}
            for name, data in items
        ]
    content = json.dumps(
        {"date": date_str, "photos": serialized},
        ensure_ascii=False,
    )
    path = f"{RECORDS_DIR}/{date_str}.json"

    try:
        try:
            existing = repo.get_contents(path)
            repo.update_file(path, f"update: {date_str}", content, existing.sha)
        except GithubException:
            repo.create_file(path, f"add: {date_str}", content)
        return True
    except Exception as e:
        st.error(f"GitHub 저장 실패: {e}")
        return False


def gh_list_dates() -> list[str]:
    """저장된 날짜 목록 (최신순)."""
    repo = _gh_repo()
    if repo is None:
        return []
    try:
        contents = repo.get_contents(RECORDS_DIR)
        dates = [
            c.name.replace(".json", "")
            for c in contents
            if c.name.endswith(".json")
        ]
        return sorted(dates, reverse=True)
    except Exception:
        return []


def gh_load(date_str: str) -> dict | None:
    """GitHub에서 해당 날짜의 사진 데이터 복원."""
    repo = _gh_repo()
    if repo is None:
        return None
    try:
        path    = f"{RECORDS_DIR}/{date_str}.json"
        content = repo.get_contents(path).decoded_content
        raw     = json.loads(content)
        photos  = {}
        for cat in FIXED_CATEGORIES:
            items        = raw["photos"].get(cat, [])
            photos[cat]  = [(it["name"], base64.b64decode(it["data"])) for it in items]
        return photos
    except Exception as e:
        st.error(f"불러오기 실패: {e}")
        return None


# ────────────────────────────────────────
# 사진 헬퍼
# ────────────────────────────────────────
def add_photos(cat: str, files):
    existing = [n for n, _ in st.session_state.photos[cat]]
    for f in files:
        data = f.read()
        name = f.name
        if name in existing:
            base, _, ext = name.rpartition(".")
            name = f"{base}_{len(existing)}.{ext}" if ext else f"{name}_{len(existing)}"
        st.session_state.photos[cat].append((name, data))
        existing.append(name)


def remove_photo(cat, idx):
    st.session_state.photos[cat].pop(idx)


def move_photo(cat, idx, d):
    lst = st.session_state.photos[cat]
    t   = idx + d
    if 0 <= t < len(lst):
        lst[idx], lst[t] = lst[t], lst[idx]


# ────────────────────────────────────────
# DOCX 생성
# ────────────────────────────────────────
def _set_cell_border(cell, color="BBBBBB"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB  = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    "4")
        tag.set(qn("w:color"), color)
        tcB.append(tag)
    tcPr.append(tcB)


def _set_cell_bg(cell, fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    tcPr.append(shd)


def _keep_with_next(paragraph):
    """다음 요소와 페이지를 함께 유지 (카테고리 제목 고립 방지)."""
    pPr = paragraph._p.get_or_add_pPr()
    kn  = OxmlElement("w:keepNext")
    pPr.append(kn)


def _fix_image(data: bytes) -> io.BytesIO:
    img = Image.open(io.BytesIO(data))
    try:
        from PIL.ExifTags import TAGS
        exif = img._getexif()
        if exif:
            for tid, val in exif.items():
                if TAGS.get(tid) == "Orientation":
                    if val == 3: img = img.rotate(180, expand=True)
                    elif val == 6: img = img.rotate(270, expand=True)
                    elif val == 8: img = img.rotate(90, expand=True)
    except Exception:
        pass
    if img.mode in ("RGBA", "P"):
        img = img.convert("RGB")
    buf = io.BytesIO()
    img.save(buf, format="JPEG", quality=85)
    buf.seek(0)
    return buf


def generate_docx(date_str: str, photos_dict: dict, per_page: int = 3) -> bytes:
    COLS = 3
    doc  = Document()
    sec  = doc.sections[0]
    sec.page_width    = Cm(21)
    sec.page_height   = Cm(29.7)
    sec.left_margin   = Cm(1.5)
    sec.right_margin  = Cm(1.5)
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(1.5)

    def add_left(text, size, bold=False, rgb=None, sb=0, sa=0, keep_next=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(sb)
        p.paragraph_format.space_after  = Pt(sa)
        if keep_next:
            _keep_with_next(p)
        r = p.add_run(text)
        r.bold       = bold
        r.font.size  = Pt(size)
        if rgb:
            r.font.color.rgb = RGBColor(*rgb)
        return p

    # ── 헤더
    add_left(DOC_TITLE, 20, bold=True, rgb=(0x1E, 0x3A, 0x5F), sa=2)
    add_left(SITE_NAME, 13, rgb=(0x2D, 0x37, 0x48), sa=1)
    add_left(f"작성일: {date_str}", 11, rgb=(0x44, 0x44, 0x44), sa=8)

    # 구분선
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after  = Pt(6)
    pPr  = p_line._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:color"), "1A56DB")
    pBdr.append(bot)
    pPr.append(pBdr)

    # ── 카테고리별 사진
    for cat in FIXED_CATEGORIES:
        photos = photos_dict.get(cat, [])
        if not photos:
            continue

        total_pages = math.ceil(len(photos) / per_page)

        for pg in range(total_pages):
            chunk  = photos[pg * per_page:(pg + 1) * per_page]
            n_rows = math.ceil(len(chunk) / COLS)

            # 카테고리 제목 — keep_next로 표와 함께 넘어가도록
            cat_label = cat if pg == 0 else f"{cat} (계속)"
            p = add_left(
                f"■ {cat_label}", 12,
                bold=True,
                rgb=(0x1A, 0x56, 0xDB),
                sb=8, sa=3,
                keep_next=True,   # ← 핵심: 아래 표와 페이지 함께 유지
            )

            # 이미지 표 (캡션 행 없음)
            table = doc.add_table(rows=n_rows, cols=COLS)
            table.style = "Table Grid"
            for row in table.rows:
                for cell in row.cells:
                    cell.width = Cm(5.9)

            for ri in range(n_rows):
                for ci in range(COLS):
                    pidx = ri * COLS + ci
                    ic   = table.rows[ri].cells[ci]
                    ic.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    _set_cell_border(ic)

                    if pidx < len(chunk):
                        _, data = chunk[pidx]
                        ip = ic.paragraphs[0]
                        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            ip.add_run().add_picture(_fix_image(data), width=Cm(5.6))
                        except Exception as e:
                            ip.add_run(f"[오류: {e}]").font.size = Pt(7)
                    else:
                        _set_cell_bg(ic, "F9FAFB")

            if pg < total_pages - 1:
                from docx.enum.text import WD_BREAK
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════
# UI
# ════════════════════════════════════════
st.markdown('<div class="main-title">🌡 온열질환 예방 현장 관리</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">카테고리별 사진을 업로드하고 저장하면 GitHub에 날짜별로 기록이 보관됩니다.</div>',
            unsafe_allow_html=True)

github_ok = _gh_repo() is not None

if not github_ok:
    st.info(
        "⚠️ GitHub 연동이 설정되지 않았습니다. "
        "`.streamlit/secrets.toml`에 `[github]` 설정을 추가하면 날짜별 저장/불러오기가 활성화됩니다.",
        icon="ℹ️",
    )

left, right = st.columns([1, 1.8], gap="large")

# ══════════════════════
# 왼쪽 패널
# ══════════════════════
with left:

    # 작성일
    st.markdown('<div class="section-hd">📅 작성일</div>', unsafe_allow_html=True)
    work_date = st.date_input(
        "작성일",
        value=st.session_state.work_date,
        label_visibility="collapsed",
    )
    st.session_state.work_date = work_date
    date_str = work_date.strftime("%Y-%m-%d")

    st.divider()

    # 사진 현황
    st.markdown('<div class="section-hd">📁 사진 현황</div>', unsafe_allow_html=True)
    total = 0
    for cat in FIXED_CATEGORIES:
        n = len(st.session_state.photos[cat])
        total += n
        st.markdown(
            f'<div class="cat-header">{cat}'
            f'<span class="count-badge">{n}장</span></div>',
            unsafe_allow_html=True,
        )
    st.write(f"전체 **{total}장**")

    st.divider()

    # 저장 & Word 생성
    st.markdown('<div class="section-hd">💾 저장 및 생성</div>', unsafe_allow_html=True)
    per_page = st.radio("페이지당 사진 수", [3, 6], horizontal=True)

    c_save, c_make = st.columns(2)

    with c_save:
        if st.button(
            "💾 GitHub 저장",
            use_container_width=True,
            disabled=(total == 0 or not github_ok),
            help="GitHub에 날짜별 JSON으로 저장합니다.",
        ):
            with st.spinner("GitHub에 저장 중..."):
                ok = gh_save(date_str, st.session_state.photos)
            if ok:
                st.success(f"✅ {date_str} 저장 완료!")

    with c_make:
        if st.button(
            "📄 Word 생성",
            type="primary",
            use_container_width=True,
            disabled=(total == 0),
        ):
            with st.spinner("생성 중..."):
                docx_bytes = generate_docx(date_str, st.session_state.photos, per_page)
            st.download_button(
                "⬇️ 다운로드",
                data=docx_bytes,
                file_name=f"온열질환예방현장관리_{date_str.replace('-','')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
            st.success("✅ 생성 완료!")

    st.divider()

    # 날짜별 기록
    st.markdown('<div class="section-hd">🗂 저장된 기록</div>', unsafe_allow_html=True)

    if not github_ok:
        st.caption("GitHub 연동 후 사용 가능합니다.")
    else:
        if st.button("🔄 목록 새로고침", use_container_width=True):
            st.session_state["saved_dates"] = gh_list_dates()

        if "saved_dates" not in st.session_state:
            st.session_state["saved_dates"] = gh_list_dates()

        saved_dates = st.session_state["saved_dates"]

        if not saved_dates:
            st.info("저장된 기록이 없습니다.")
        else:
            for d in saved_dates:
                col_d, col_l = st.columns([2, 1])
                col_d.write(f"📅 {d}")
                if col_l.button("불러오기", key=f"load_{d}", use_container_width=True):
                    with st.spinner(f"{d} 불러오는 중..."):
                        loaded = gh_load(d)
                    if loaded:
                        st.session_state.photos    = loaded
                        st.session_state.work_date = datetime.strptime(d, "%Y-%m-%d").date()
                        st.success(f"✅ {d} 불러오기 완료!")
                        st.rerun()

    st.write("")
    if st.button("🗑 현재 사진 모두 지우기", use_container_width=True):
        st.session_state.photos = {cat: [] for cat in FIXED_CATEGORIES}
        st.rerun()


# ────────────────────────────────────────
# 클립보드 붙여넣기 컴포넌트
# ────────────────────────────────────────
def paste_zone(cat: str):
    """
    Ctrl+V 붙여넣기 영역.
    이미지를 base64로 읽어 숨겨진 st.text_area에 기록 → Streamlit이 감지해서 세션에 추가.
    """
    safe_cat = cat.replace("/", "_").replace(" ", "_")
    area_key = f"paste_data_{safe_cat}"

    # 숨겨진 텍스트에어리어 (base64 데이터 수신용)
    pasted_b64 = st.text_area(
        label=area_key,
        key=area_key,
        label_visibility="collapsed",
        height=1,
        placeholder="",
    )

    # 붙여넣기 UI (iframe)
    uid = safe_cat
    html = f"""
<style>
  #pzone-{uid} {{
    border: 2px dashed #93C5FD;
    border-radius: 10px;
    padding: 18px;
    text-align: center;
    font-family: sans-serif;
    font-size: 14px;
    color: #3B82F6;
    background: #EFF6FF;
    cursor: pointer;
    outline: none;
    min-height: 60px;
    display: flex; align-items: center; justify-content: center;
    transition: background 0.15s;
    user-select: none;
  }}
  #pzone-{uid}.active {{ background:#DBEAFE; border-color:#3B82F6; }}
  #pzone-{uid}.ok     {{ background:#ECFDF5; border-color:#10B981; color:#065F46; }}
</style>
<div id="pzone-{uid}" tabindex="0">
  📋 여기를 클릭 후 <b>Ctrl+V</b> 로 사진 붙여넣기
</div>
<script>
(function() {{
  const zone = document.getElementById('pzone-{uid}');

  zone.addEventListener('click', () => zone.focus());
  zone.addEventListener('focus', () => zone.classList.add('active'));
  zone.addEventListener('blur',  () => zone.classList.remove('active'));

  function handlePaste(e) {{
    const items = (e.clipboardData || e.originalEvent.clipboardData).items;
    for (let i = 0; i < items.length; i++) {{
      if (items[i].type.startsWith('image/')) {{
        e.preventDefault();
        const blob = items[i].getAsFile();
        const reader = new FileReader();
        reader.onload = function(ev) {{
          const b64 = ev.target.result;   // data:image/png;base64,...

          // 부모 Streamlit 페이지의 textarea에 값 주입 후 change 이벤트 발생
          const parent = window.parent.document;
          const textareas = parent.querySelectorAll('textarea');
          let target = null;
          textareas.forEach(ta => {{
            if (ta.getAttribute('aria-label') === '{area_key}' ||
                ta.closest('[data-testid]')?.querySelector('label')?.textContent?.trim() === '{area_key}') {{
              target = ta;
            }}
          }});
          // fallback: placeholder로 찾기
          if (!target) {{
            textareas.forEach(ta => {{ if (ta.placeholder === '') target = ta; }});
          }}
          if (target) {{
            const nativeInputValueSetter = Object.getOwnPropertyDescriptor(
              window.parent.HTMLTextAreaElement.prototype, 'value').set;
            nativeInputValueSetter.call(target, b64);
            target.dispatchEvent(new Event('input', {{ bubbles: true }}));
            target.dispatchEvent(new Event('change', {{ bubbles: true }}));
          }}

          zone.innerHTML = '✅ 추가됐습니다! 계속 붙여넣거나 왼쪽에서 Word 생성하세요.';
          zone.classList.add('ok');
          setTimeout(() => {{
            zone.innerHTML = '📋 여기를 클릭 후 <b>Ctrl+V</b> 로 사진 붙여넣기';
            zone.classList.remove('ok');
          }}, 2500);
        }};
        reader.readAsDataURL(blob);
        break;
      }}
    }}
  }}

  zone.addEventListener('paste', handlePaste);
  // 탭 내 어디서나 붙여넣기 가능하도록 부모 문서에도 등록
  try {{
    window.parent.document.addEventListener('paste', handlePaste);
  }} catch(e) {{}}
}})();
</script>
"""
    components.html(html, height=90, scrolling=False)

    # text_area에 base64가 들어왔으면 bytes로 변환해서 반환
    if pasted_b64 and pasted_b64.startswith("data:image"):
        try:
            _, b64 = pasted_b64.split(",", 1)
            img_bytes = base64.b64decode(b64)
            # 읽은 뒤 초기화 (다음 리런에서 중복 방지)
            st.session_state[area_key] = ""
            return img_bytes
        except Exception:
            pass
    return None


# ══════════════════════
# 오른쪽 패널
# ══════════════════════
with right:
    st.markdown('<div class="section-hd">🖼 사진 추가</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="tip">'
        '📋 <b>Ctrl+V</b> 붙여넣기 또는 파일 선택 모두 가능합니다. '
        '↑↓ 버튼으로 순서를 조정하세요.'
        '</div>',
        unsafe_allow_html=True,
    )

    tab_labels = [f"{icon} {cat}" for icon, cat in zip(TAB_ICONS, FIXED_CATEGORIES)]
    tabs       = st.tabs(tab_labels)

    for tab, cat in zip(tabs, FIXED_CATEGORIES):
        with tab:

            # ── Ctrl+V 붙여넣기 영역
            img_bytes = paste_zone(cat)
            if img_bytes:
                safe_cat = cat.replace("/", "_").replace(" ", "_")
                fname = f"paste_{safe_cat}_{int(time.time())}.png"
                st.session_state.photos[cat].append((fname, img_bytes))
                st.rerun()

            # ── 파일 업로더 (기존)
            uploaded = st.file_uploader(
                "또는 파일 선택",
                type=["jpg", "jpeg", "png", "bmp", "webp"],
                accept_multiple_files=True,
                key=f"up_{cat}",
                label_visibility="visible",
            )
            if uploaded:
                add_photos(cat, uploaded)
                st.rerun()

            # ── 사진 미리보기
            photos = st.session_state.photos[cat]
            if not photos:
                st.info("아직 사진이 없습니다.")
            else:
                st.write(f"총 **{len(photos)}장**")
                PCOLS = 3
                for ri in range(math.ceil(len(photos) / PCOLS)):
                    cols = st.columns(PCOLS)
                    for ci in range(PCOLS):
                        pidx = ri * PCOLS + ci
                        if pidx < len(photos):
                            fname, data = photos[pidx]
                            with cols[ci]:
                                try:
                                    st.image(data, use_container_width=True)
                                except Exception:
                                    st.write("⚠ 미리보기 불가")
                                b1, b2, b3 = st.columns(3)
                                with b1:
                                    if st.button("↑", key=f"u_{cat}_{pidx}"):
                                        move_photo(cat, pidx, -1); st.rerun()
                                with b2:
                                    if st.button("↓", key=f"d_{cat}_{pidx}"):
                                        move_photo(cat, pidx,  1); st.rerun()
                                with b3:
                                    if st.button("🗑", key=f"x_{cat}_{pidx}"):
                                        remove_photo(cat, pidx); st.rerun()
